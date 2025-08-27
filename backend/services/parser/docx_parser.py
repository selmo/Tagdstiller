from pathlib import Path
from typing import Optional
from .base import DocumentParser, ParseResult, DocumentMetadata
from utils.text_cleaner import TextCleaner

class DocxParser(DocumentParser):
    """DOCX 파일 파서 (python-docx 사용)"""
    
    def __init__(self):
        super().__init__("docx_parser")
        self.supported_extensions = ['.docx', '.docm']
        self.supported_mime_types = [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/vnd.ms-word.document.macroEnabled.12'
        ]
    
    def parse(self, file_path: Path) -> ParseResult:
        """DOCX 파일을 파싱합니다."""
        try:
            # python-docx 임포트 확인
            try:
                from docx import Document
                from docx.opc.exceptions import PackageNotFoundError
            except ImportError:
                return self.create_error_result(
                    "python-docx 라이브러리가 설치되지 않았습니다. 'pip install python-docx' 실행이 필요합니다.",
                    file_path
                )
            
            # 파일 기본 정보
            file_info = self.get_file_info(file_path)
            
            # DOCX 문서 열기
            try:
                doc = Document(str(file_path))
            except PackageNotFoundError:
                return self.create_error_result(
                    "유효하지 않은 DOCX 파일입니다.",
                    file_path
                )
            
            # 문서 텍스트 추출
            text_parts = []
            
            # 본문 텍스트 추출
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # 단락별 텍스트 정제
                    cleaned_text = TextCleaner.clean_text(paragraph.text)
                    if cleaned_text.strip():
                        text_parts.append(cleaned_text)
            
            # 표 텍스트 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            # 셀별 텍스트 정제
                            cleaned_cell_text = TextCleaner.clean_text(cell.text.strip())
                            if cleaned_cell_text.strip():
                                row_text.append(cleaned_cell_text)
                    if row_text:
                        text_parts.append('\t'.join(row_text))
            
            # 전체 텍스트 결합
            full_text = '\n'.join(text_parts)
            word_count = len(full_text.split()) if full_text else 0
            
            # 문서 속성에서 메타데이터 추출
            core_props = doc.core_properties
            
            # 확장된 메타데이터 생성
            from datetime import datetime
            import os
            
            # 앱 버전 가져오기
            app_version = os.getenv("APP_VERSION", "1.0.0")
            
            # 언어 감지
            def detect_language(text_sample):
                if not text_sample:
                    return "unknown"
                korean_chars = sum(1 for c in text_sample[:1000] if '\uAC00' <= c <= '\uD7A3')
                if korean_chars > len(text_sample[:1000]) * 0.1:
                    return "ko"
                english_chars = sum(1 for c in text_sample[:1000] if c.isascii() and c.isalpha())
                if english_chars > len(text_sample[:1000]) * 0.5:
                    return "en"
                return "unknown"
                
            detected_language = detect_language(full_text)
            
            # 페이지 수 추정 (평균 250단어/페이지 기준)
            estimated_pages = max(1, word_count // 250)
            
            # Dublin Core 메타데이터 매핑
            dc_title = core_props.title or file_path.stem
            dc_creator = core_props.author
            dc_date = str(core_props.created) if core_props.created else None
            
            metadata = DocumentMetadata(
                # 기존 필드 (호환성)
                title=dc_title,
                author=dc_creator,
                created_date=dc_date,
                modified_date=str(core_props.modified) if core_props.modified else None,
                word_count=word_count,
                file_size=file_info.get("file_size"),
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                
                # Dublin Core 메타데이터
                dc_title=dc_title,
                dc_creator=dc_creator,
                dc_subject=core_props.subject,
                dc_description=core_props.comments or f"Word 문서, 약 {estimated_pages}페이지",
                dc_publisher=None,
                dc_contributor=core_props.last_modified_by,
                dc_date=dc_date,
                dc_type="document",
                dc_format="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                dc_identifier=file_path.name,
                dc_source=str(file_path),
                dc_language=detected_language,
                dc_rights=None,
                
                # Dublin Core Terms
                dcterms_created=dc_date,
                dcterms_modified=str(core_props.modified) if core_props.modified else None,
                dcterms_extent=f"{file_info.get('file_size', 0)} bytes",
                dcterms_medium="digital",
                
                # 파일 메타데이터
                file_name=file_path.name,
                file_path=str(file_path),
                file_extension=file_path.suffix.lower(),
                
                # 문서 메타데이터
                doc_page_count=estimated_pages,
                doc_word_count=word_count,
                doc_character_count=len(full_text) if full_text else 0,
                doc_type_code="docx",
                doc_supported="yes",
                
                # 애플리케이션 메타데이터
                app_version=app_version,
                
                # 파서 정보
                parser_name=self.parser_name,
                parser_version="1.0"
            )
            
            return ParseResult(
                text=full_text,
                metadata=metadata,
                success=True,
                parser_name=self.parser_name
            )
            
        except Exception as e:
            return self.create_error_result(f"DOCX 파싱 오류: {str(e)}", file_path)
    
    def extract_styles_info(self, file_path: Path) -> dict:
        """DOCX 파일의 스타일 정보를 추출합니다."""
        try:
            from docx import Document
            doc = Document(str(file_path))
            
            styles_info = {
                "paragraph_styles": [],
                "character_styles": [],
                "table_styles": [],
                "numbering_styles": []
            }
            
            for style in doc.styles:
                style_info = {
                    "name": style.name,
                    "type": str(style.type),
                    "builtin": style.builtin
                }
                
                if style.type.name == 'PARAGRAPH':
                    styles_info["paragraph_styles"].append(style_info)
                elif style.type.name == 'CHARACTER':
                    styles_info["character_styles"].append(style_info)
                elif style.type.name == 'TABLE':
                    styles_info["table_styles"].append(style_info)
                elif style.type.name == 'NUMBERING':
                    styles_info["numbering_styles"].append(style_info)
            
            return styles_info
            
        except Exception as e:
            return {"error": str(e)}
    
    def get_document_structure(self, file_path: Path) -> dict:
        """DOCX 문서의 구조 정보를 반환합니다."""
        try:
            from docx import Document
            doc = Document(str(file_path))
            
            structure = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "section_count": len(doc.sections),
                "has_headers": False,
                "has_footers": False,
                "has_images": False
            }
            
            # 헤더/푸터 확인
            for section in doc.sections:
                if section.header.paragraphs:
                    structure["has_headers"] = True
                if section.footer.paragraphs:
                    structure["has_footers"] = True
            
            # 이미지 확인 (간단한 방법)
            for paragraph in doc.paragraphs:
                if paragraph._element.xpath('.//pic:pic'):
                    structure["has_images"] = True
                    break
            
            return structure
            
        except Exception as e:
            return {"error": str(e)}